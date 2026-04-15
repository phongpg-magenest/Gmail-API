from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(0x0f, 0x3c, 0x78)
    return p

def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Background via shading on paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    run = p.add_run(code)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2d, 0x2d, 0x2d)
    return p

def add_note(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.3)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'FFF8E1')
    pPr.append(shd)
    run = p.add_run('Lưu ý: ')
    run.bold = True
    run.font.color.rgb = RGBColor(0xf5, 0x7c, 0x00)
    run2 = p.add_run(text)
    run2.font.color.rgb = RGBColor(0x4a, 0x4a, 0x4a)

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        set_cell_bg(cell, '1F4E79')
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = 'FFFFFF' if ri % 2 == 0 else 'EBF5FB'
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            # Handle inline code (backtick)
            parts = re.split(r'`([^`]+)`', val)
            for pi, part in enumerate(parts):
                if pi % 2 == 1:
                    r = p.add_run(part)
                    r.font.name = 'Courier New'
                    r.font.size = Pt(9)
                    r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
                else:
                    if part:
                        r = p.add_run(part)
                        r.font.size = Pt(10)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()

def add_para(doc, text):
    """Add paragraph with basic bold (**text**) and inline code (`text`) support."""
    p = doc.add_paragraph()
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        else:
            if part:
                p.add_run(part)
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.3)
    parts = re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        elif part.startswith('`') and part.endswith('`'):
            r = p.add_run(part[1:-1])
            r.font.name = 'Courier New'
            r.font.size = Pt(10)
            r.font.color.rgb = RGBColor(0xc0, 0x39, 0x2b)
        else:
            if part:
                p.add_run(part)

# ============================================================
# DOCUMENT CONTENT
# ============================================================

add_heading(doc, 'Gmail API — Tài liệu tích hợp', 1)
doc.add_paragraph()

# ---- Section 1 ----
add_heading(doc, '1. Setup Google Cloud', 2)

add_heading(doc, '1.1 Tạo project trên Google Cloud Console', 3)
add_bullet(doc, 'Truy cập console.cloud.google.com')
add_bullet(doc, 'Tạo project mới (hoặc dùng project có sẵn)')
add_bullet(doc, 'Vào **APIs & Services → Library**, tìm và bật **Gmail API**')
doc.add_paragraph()

add_heading(doc, '1.2 Tạo OAuth 2.0 Credentials', 3)
add_bullet(doc, 'Vào **APIs & Services → Credentials → Create Credentials → OAuth client ID**')
add_bullet(doc, 'Application type: **Web application**')
add_bullet(doc, 'Thêm Authorized redirect URIs:')
add_code_block(doc, 'http://localhost:3000/auth/google/callback')
add_bullet(doc, 'Lưu lại `Client ID` và `Client Secret`')
doc.add_paragraph()

add_heading(doc, '1.3 Cấu hình OAuth Consent Screen', 3)
add_bullet(doc, 'Vào **APIs & Services → OAuth consent screen**')
add_bullet(doc, 'User Type: **External** (hoặc Internal nếu dùng Google Workspace)')
add_bullet(doc, 'Điền App name, support email')
add_bullet(doc, 'Thêm scopes:')
add_code_block(doc, 'https://www.googleapis.com/auth/gmail.readonly')
add_bullet(doc, 'Thêm test users (email của tài khoản Gmail muốn kết nối)')
doc.add_paragraph()
add_note(doc, 'Khi app ở trạng thái "Testing", chỉ các email trong danh sách test users mới đăng nhập được. Để mở rộng cần submit Google verification.')
doc.add_paragraph()

add_heading(doc, '1.4 Cấu hình .env', 3)
add_code_block(doc, (
    'GOOGLE_CLIENT_ID=your_client_id_here.apps.googleusercontent.com\n'
    'GOOGLE_CLIENT_SECRET=your_client_secret_here\n'
    'TOKEN_ENCRYPTION_SECRET=random_32_char_string_for_aes_encryption\n'
    'GEMINI_API_KEY=your_gemini_api_key\n'
    'DEMO_USER_EMAIL=your@email.com\n'
    'DEMO_USER_PASSWORD=yourpassword'
))
doc.add_paragraph()

# ---- Section 2 ----
add_heading(doc, '2. OAuth Flow', 2)
add_code_block(doc, (
    'User → GET /auth/google/start\n'
    '     → redirect đến Google OAuth consent screen\n'
    '     → User chấp thuận\n'
    '     → Google redirect về /auth/google/callback?code=...\n'
    '     → App dùng code đổi lấy access_token + refresh_token\n'
    '     → Lưu refresh_token (đã mã hóa AES) vào SQLite\n'
    '     → Dùng refresh_token để lấy access_token mỗi khi gọi Gmail API'
))
doc.add_paragraph()
add_para(doc, '**Scopes sử dụng:**')
add_table(doc,
    ['Scope', 'Quyền'],
    [
        ['`gmail.readonly`', 'Đọc mail, không thể gửi/xóa'],
        ['`userinfo.email`', 'Lấy email của tài khoản'],
        ['`userinfo.profile`', 'Lấy thông tin cơ bản'],
    ],
    [2.5, 4.0]
)

# ---- Section 3 ----
add_heading(doc, '3. Gmail API — Threads', 2)
add_para(doc, 'Gmail tổ chức email theo **thread** (luồng hội thoại). Mỗi thread gồm nhiều message cùng subject/conversation.')
doc.add_paragraph()

add_heading(doc, '3.1 List Threads', 3)
add_code_block(doc, 'GET https://gmail.googleapis.com/gmail/v1/users/me/threads')
doc.add_paragraph()
add_para(doc, '**Query params:**')
add_table(doc,
    ['Param', 'Mô tả', 'Ví dụ'],
    [
        ['`q`', 'Gmail search query', '`from:@magenest.com newer_than:7d`'],
        ['`labelIds`', 'Lọc theo label', '`INBOX`, `SPAM`, `CATEGORY_PROMOTIONS`'],
        ['`maxResults`', 'Số thread tối đa', '`20`'],
        ['`pageToken`', 'Phân trang', 'token từ response trước'],
    ],
    [1.5, 2.5, 2.5]
)
add_para(doc, '**Response JSON:**')
add_code_block(doc, (
    '{\n'
    '  "threads": [\n'
    '    { "id": "18f3a1b2c3d4e5f6", "snippet": "Xin chào, tôi muốn hỏi về..." },\n'
    '    { "id": "18f3a1b2c3d4e5f7", "snippet": "Re: Báo giá dự án..." }\n'
    '  ],\n'
    '  "nextPageToken": "abc123",\n'
    '  "resultSizeEstimate": 42\n'
    '}'
))
doc.add_paragraph()

add_heading(doc, '3.2 Get Thread Detail', 3)
add_code_block(doc, 'GET https://gmail.googleapis.com/gmail/v1/users/me/threads/{threadId}')
doc.add_paragraph()
add_para(doc, '**Query params:**')
add_table(doc,
    ['Param', 'Mô tả'],
    [
        ['`format`', '`full` (body đầy đủ), `metadata` (chỉ headers + snippet), `minimal`'],
        ['`metadataHeaders`', 'Headers cần lấy khi dùng metadata format, vd: ["From","To","Subject","Date"]'],
    ],
    [2.0, 4.5]
)
add_para(doc, '**Response JSON (format: metadata) — ví dụ thread 3 tin:**')
add_code_block(doc, (
    '{\n'
    '  "id": "18f3a1b2c3d4e5f6",\n'
    '  "historyId": "1234567",\n'
    '  "messages": [\n'
    '    {\n'
    '      "id": "18f3a1b2c3d4e5f6",\n'
    '      "threadId": "18f3a1b2c3d4e5f6",\n'
    '      "labelIds": ["INBOX", "UNREAD", "IMPORTANT"],\n'
    '      "snippet": "Xin chào anh Phong, tôi muốn hỏi về gói dịch vụ...",\n'
    '      "payload": {\n'
    '        "headers": [\n'
    '          { "name": "From",    "value": "Nguyen Van A <nguyenvana@example.com>" },\n'
    '          { "name": "To",      "value": "phongpg@magenest.com" },\n'
    '          { "name": "Subject", "value": "Hỏi về gói dịch vụ Magento" },\n'
    '          { "name": "Date",    "value": "Mon, 13 Apr 2026 09:15:00 +0700" }\n'
    '        ]\n'
    '      }\n'
    '    },\n'
    '    {\n'
    '      "id": "18f3a1b2c3d4e5f8",\n'
    '      "threadId": "18f3a1b2c3d4e5f6",\n'
    '      "labelIds": ["SENT"],\n'
    '      "snippet": "Chào anh, cảm ơn anh đã liên hệ. Bên mình có các gói...",\n'
    '      "payload": {\n'
    '        "headers": [\n'
    '          { "name": "From",    "value": "Phong <phongpg@magenest.com>" },\n'
    '          { "name": "To",      "value": "nguyenvana@example.com" },\n'
    '          { "name": "Subject", "value": "Re: Hỏi về gói dịch vụ Magento" },\n'
    '          { "name": "Date",    "value": "Mon, 13 Apr 2026 10:30:00 +0700" }\n'
    '        ]\n'
    '      }\n'
    '    },\n'
    '    {\n'
    '      "id": "18f3a1b2c3d4e5fa",\n'
    '      "threadId": "18f3a1b2c3d4e5f6",\n'
    '      "labelIds": ["INBOX"],\n'
    '      "snippet": "Cảm ơn anh rất nhiều, anh cho mình xin báo giá chi tiết không?",\n'
    '      "payload": {\n'
    '        "headers": [\n'
    '          { "name": "From",    "value": "Nguyen Van A <nguyenvana@example.com>" },\n'
    '          { "name": "To",      "value": "phongpg@magenest.com" },\n'
    '          { "name": "Subject", "value": "Re: Hỏi về gói dịch vụ Magento" },\n'
    '          { "name": "Date",    "value": "Mon, 13 Apr 2026 14:05:00 +0700" }\n'
    '        ]\n'
    '      }\n'
    '    }\n'
    '  ]\n'
    '}'
))
doc.add_paragraph()

# ---- Section 4 ----
add_heading(doc, '4. Label IDs', 2)
add_para(doc, 'Gmail dùng label để phân loại email. Một message có thể có nhiều label cùng lúc.')
doc.add_paragraph()
add_para(doc, '**System Labels:**')
add_table(doc,
    ['Label ID', 'Ý nghĩa'],
    [
        ['`INBOX`', 'Hộp thư đến'],
        ['`SENT`', 'Đã gửi'],
        ['`DRAFT`', 'Nháp'],
        ['`TRASH`', 'Thùng rác'],
        ['`SPAM`', 'Thư rác'],
        ['`STARRED`', 'Có gắn sao'],
        ['`IMPORTANT`', 'Google đánh dấu quan trọng'],
        ['`UNREAD`', 'Chưa đọc'],
    ],
    [2.5, 4.0]
)
add_para(doc, '**Category Labels (Gmail tự phân loại):**')
add_table(doc,
    ['Label ID', 'Tab trong Gmail'],
    [
        ['`CATEGORY_PERSONAL`', 'Cá nhân'],
        ['`CATEGORY_SOCIAL`', 'Mạng xã hội'],
        ['`CATEGORY_PROMOTIONS`', 'Quảng cáo'],
        ['`CATEGORY_UPDATES`', 'Thông báo'],
        ['`CATEGORY_FORUMS`', 'Diễn đàn'],
    ],
    [2.5, 4.0]
)

# ---- Section 5 ----
add_heading(doc, '5. Gmail Search Query (tham số q)', 2)
add_para(doc, 'Cú pháp tương tự thanh tìm kiếm trong Gmail:')
add_table(doc,
    ['Query', 'Ý nghĩa'],
    [
        ['`from:@magenest.com`', 'Từ domain magenest.com'],
        ['`to:phong@example.com`', 'Gửi đến email cụ thể'],
        ['`subject:báo giá`', 'Tiêu đề chứa "báo giá"'],
        ['`newer_than:7d`', 'Trong 7 ngày qua'],
        ['`newer_than:1d`', 'Trong 24 giờ qua'],
        ['`older_than:30d`', 'Cũ hơn 30 ngày'],
        ['`has:attachment`', 'Có đính kèm'],
        ['`is:unread`', 'Chưa đọc'],
        ['`is:starred`', 'Có gắn sao'],
        ['`label:INBOX`', 'Trong inbox'],
    ],
    [2.5, 4.0]
)
add_para(doc, 'Có thể kết hợp nhiều điều kiện:')
add_code_block(doc, 'from:@magenest.com newer_than:7d is:unread')
doc.add_paragraph()

# ---- Section 6 ----
add_heading(doc, '6. ThreadSummary — Cấu trúc dữ liệu trong app (custom, không phải Gmail SDK)', 2)
add_para(doc, (
    'Đây là TypeScript type **tự định nghĩa trong app**, không phải từ Gmail SDK hay API. '
    'Gmail API trả về raw JSON với headers dạng array lồng nhau (xem phần 3.2). '
    'App normalize lại thành ThreadSummary để tiện dùng trong code — '
    'thay vì phải viết messages[0].payload.headers.find(h => h.name === "From") '
    'thì chỉ cần dùng thread.firstFrom.'
))
add_code_block(doc, (
    'type ThreadSummary = {\n'
    '  id: string           // threadId (dùng làm cache key)\n'
    '  subject: string      // subject của message đầu tiên\n'
    '  messageCount: number // tổng số message trong thread\n'
    '\n'
    '  // Message đầu tiên (người khởi tạo)\n'
    '  firstFrom: string    // "Nguyen Van A <nguyenvana@example.com>"\n'
    '  firstDate: string    // "Mon, 13 Apr 2026 09:15:00 +0700"\n'
    '\n'
    '  // Message cuối cùng (trạng thái hiện tại)\n'
    '  lastFrom: string     // người gửi gần nhất\n'
    '  lastDate: string     // thời gian gần nhất\n'
    '  lastSnippet: string  // preview nội dung gần nhất\n'
    '  lastLabelIds: string[] // labels của message cuối\n'
    '\n'
    '  // Toàn bộ snippets để LLM đánh giá\n'
    '  allSnippets: string[]\n'
    '}'
))
doc.add_paragraph()

# ---- Section 7 ----
add_heading(doc, '7. LLM Classification Cache', 2)
add_para(doc, 'Kết quả phân loại được lưu vào SQLite để tránh gọi lại LLM mỗi lần load:')
add_code_block(doc, (
    'CREATE TABLE thread_classifications (\n'
    '  thread_id     TEXT PRIMARY KEY,\n'
    '  important     INTEGER NOT NULL,  -- 0 hoặc 1\n'
    '  reason        TEXT NOT NULL,     -- lý do ngắn gọn từ LLM\n'
    '  message_count INTEGER NOT NULL,  -- dùng để detect khi có reply mới\n'
    '  classified_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP\n'
    ');'
))
doc.add_paragraph()
add_para(doc, '**Logic invalidation cache:**')
add_table(doc,
    ['Trường hợp', 'Hành động'],
    [
        ['thread_id có trong DB && message_count không đổi', 'Dùng cache SQLite, không gọi LLM'],
        ['thread_id có trong DB && message_count tăng', 'Có reply mới → gọi LLM lại, cập nhật DB'],
        ['thread_id chưa có trong DB', 'Thread mới → gọi LLM, lưu kết quả vào DB'],
    ],
    [3.5, 3.0]
)

# Save
out = '/home/phong/PROJECT/Mail/docs/gmail-api.docx'
doc.save(out)
print(f'Saved: {out}')
